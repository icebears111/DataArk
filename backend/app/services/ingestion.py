"""
文档导入服务。

负责：
1. 解析 PDF/Word/Markdown/TXT 文件
2. 文本分块（chunk）
3. 向量化（embedding）
4. 存入 ChromaDB 向量数据库
"""

import os
import uuid
from datetime import datetime
from typing import List, Optional

from langchain_classic.schema import Document as LCDocument
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter

from app.config import settings

# 上传文件存储目录
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "uploads")


def _ensure_upload_dir():
    """确保上传目录存在。"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)


def _get_embeddings():
    """
    获取 embedding 模型。
    
    注意：LangChain 的 OpenAIEmbeddings 会把文本转成 token ID 再发请求，
    但 DashScope 的 embedding API 只接受原始字符串。
    
    所以这里不使用 LangChain 的 OpenAIEmbeddings，而是用 openai 库直接调用。
    """
    if settings.LLM_PROVIDER == "ollama":
        from langchain_ollama import OllamaEmbeddings
        return OllamaEmbeddings(model="qwen2.5:7b", base_url=settings.OLLAMA_BASE_URL)

    from langchain_core.embeddings import Embeddings as BaseEmbeddings

    class DashScopeEmbeddings(BaseEmbeddings):
        """
        自定义 Embeddings 类，适配 DashScope 的 embedding API。
        
        和 OpenAIEmbeddings 的区别：
        - OpenAIEmbeddings：把文本转成 token ID 再发送（DashScope 不支持）
        - DashScopeEmbeddings：直接发送原始字符串
        """

        def __init__(self):
            from openai import OpenAI
            self.client = OpenAI(
                api_key=settings.LLM_API_KEY,
                base_url=settings.EMBEDDING_API_BASE,
            )
            self.model = settings.EMBEDDING_MODEL

        def embed_documents(self, texts):
            """批量向量化文档文本。"""
            response = self.client.embeddings.create(
                model=self.model,
                input=texts,
            )
            return [item.embedding for item in response.data]

        def embed_query(self, text):
            """向量化单个查询文本。"""
            response = self.client.embeddings.create(
                model=self.model,
                input=text,
            )
            return response.data[0].embedding

    return DashScopeEmbeddings()


def _get_vector_store(collection_name: str = "documents"):
    """
    获取 ChromaDB 向量数据库实例。
    
    ChromaDB 是一个本地向量数据库，用来存文档的向量。
    每个文档被切成块，每块变成一个向量存入 ChromaDB。
    搜索时把问题也向量化，在 ChromaDB 里找最相似的块。
    """
    from langchain_chroma import Chroma

    return Chroma(
        collection_name=collection_name,
        embedding_function=_get_embeddings(),
        persist_directory=settings.CHROMA_PERSIST_DIR,
    )


def parse_document(file_path: str) -> str:
    """
    解析文档，提取纯文本。
    
    支持格式：PDF、Markdown、TXT
    注意：Word (.docx) 需要额外安装 python-docx 库。
    """
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif file_ext in (".md", ".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    elif file_ext == ".docx":
        # docx 需要 python-docx 库
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return "Word 文档解析需要安装 python-docx 库"

    else:
        raise ValueError(f"不支持的文件格式: {file_ext}")


def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    把长文本切成小块。
    
    为什么需要切块？
    1. LLM 有上下文长度限制
    2. 小块更容易精确匹配用户问题
    3. 可以返回多个相关块增加信息量
    
    chunk_size=500：每块约 500 字符
    chunk_overlap=50：相邻块重叠 50 字符，避免信息断裂
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", ".", " ", ""],
    )
    return splitter.split_text(text)


def index_document(file_path: str, doc_id: int) -> int:
    """
    完整的文档索引流程。
    
    1. 解析 → 2. 分块 → 3. 向量化 → 4. 存入 ChromaDB
    
    返回：索引的文本块数量
    """
    # 1. 解析文档
    text = parse_document(file_path)
    if not text.strip():
        raise ValueError("文档内容为空")

    # 2. 分块
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("文档无法分块")

    # 3. 创建 LangChain Document 对象（附带元数据）
    lc_docs = []
    for i, chunk in enumerate(chunks):
        lc_docs.append(LCDocument(
            page_content=chunk,
            metadata={
                "doc_id": doc_id,
                "chunk_index": i,
                "file_path": file_path,
            },
        ))

    # 4. 存入 ChromaDB
    vector_store = _get_vector_store()
    vector_store.add_documents(lc_docs)

    return len(chunks)


def search_documents(query: str, k: int = 5) -> List[dict]:
    """
    搜索文档。
    
    流程：
    1. 把问题向量化
    2. 在 ChromaDB 中找最相似的 k 个块
    3. 返回匹配结果
    """
    vector_store = _get_vector_store()
    results = vector_store.similarity_search_with_score(query, k=k)

    docs = []
    for doc, score in results:
        docs.append({
            "content": doc.page_content,
            "metadata": doc.metadata,
            "score": round(score, 4),
        })
    return docs


def delete_document_index(doc_id: int):
    """
    从 ChromaDB 中删除某个文档的所有向量块。
    根据 doc_id 从 metadata 中匹配删除。
    """
    vector_store = _get_vector_store()
    # 获取所有 doc_id 匹配的块
    all_docs = vector_store.get()
    if not all_docs or "metadatas" not in all_docs:
        return

    ids_to_delete = []
    for i, meta in enumerate(all_docs["metadatas"]):
        if meta and meta.get("doc_id") == doc_id:
            ids_to_delete.append(all_docs["ids"][i])

    if ids_to_delete:
        vector_store.delete(ids_to_delete)
